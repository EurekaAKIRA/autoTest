from docx import Document
from app.models.schemas import TestCase
import logging
import traceback
import io

# 配置日志记录器
logger = logging.getLogger(__name__)

class ExportService:
    async def export_to_word(self, test_cases: list[TestCase]) -> bytes:
        """
        导出测试用例为Word文档
        """
        try:
            logger.info("开始导出Word文档")
            doc = Document()
            doc.add_heading("测试用例集", 0)
            
            for idx, case in enumerate(test_cases, 1):
                doc.add_heading(f"用例 {idx}: {case.title}", level=1)
                doc.add_paragraph(f"【描述】{case.description}")
                doc.add_paragraph("【步骤】")
                for i, step in enumerate(case.steps, 1):
                    doc.add_paragraph(f"{i}. {step}", style="List Number")
                doc.add_paragraph(f"【预期结果】{case.expected_result}")
                doc.add_paragraph("")  # 空行分隔
                
            # 保存到内存中
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            
            logger.info("Word文档导出成功")
            return file_stream.getvalue()
        except Exception as e:
            logger.error(f"导出Word文档失败: {str(e)}")
            logger.error(traceback.format_exc())
            raise Exception(f"导出Word文档失败: {str(e)}")
            
    async def export_to_script(self, test_cases: list[TestCase], script_type: str = "pytest") -> str:
        """
        导出测试用例为Python脚本
        """
        try:
            logger.info("开始导出Python脚本")
            script_lines = []
            
            if script_type == "pytest":
                script_lines.append("import pytest\n")
            elif script_type == "unittest":
                script_lines.append("import unittest\n")
            
            for idx, case in enumerate(test_cases, 1):
                if script_type == "pytest":
                    script_lines.append(f"def test_case_{idx}():")
                    script_lines.append(f"    \"\"\"{case.title}: {case.description}\"\"\"")
                    for step in case.steps:
                        script_lines.append(f"    # {step}")
                    script_lines.append(f"    # 预期结果: {case.expected_result}")
                    script_lines.append("")
                elif script_type == "unittest":
                    script_lines.append(f"class TestCase{idx}(unittest.TestCase):")
                    script_lines.append(f"    def test_{idx}(self):")
                    script_lines.append(f"        \"\"\"{case.title}: {case.description}\"\"\"")
                    for step in case.steps:
                        script_lines.append(f"        # {step}")
                    script_lines.append(f"        # 预期结果: {case.expected_result}")
                    script_lines.append("")
            
            if script_type == "unittest":
                script_lines.append("if __name__ == '__main__':")
                script_lines.append("    unittest.main()")
            
            script_content = "\n".join(script_lines)
            logger.info("Python脚本导出成功")
            return script_content
        except Exception as e:
            logger.error(f"导出Python脚本失败: {str(e)}")
            logger.error(traceback.format_exc())
            raise Exception(f"导出Python脚本失败: {str(e)}") 